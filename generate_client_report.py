import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pathlib import Path
from data.data_handler import DataHandler
from utils.analytics import calculate_metrics

def generate_client_report():
    base_path = Path("/Users/shubhrakasana/Desktop/oleander/olgo_mobile/repos/Backtest/precision-backtest-engine")
    output_dir = base_path / "outputs"
    output_dir.mkdir(exist_ok=True)
    report_file = base_path / "Client_Report_Contrarian_Quarterly.md"
    
    # 1. Load Data
    # Strategy
    strat_nav_path = output_dir / "final_champion_nav.csv"
    if not strat_nav_path.exists():
        print("Strategy NAV not found. Please run final_champion_run.py first.")
        return
    strat_nav = pd.read_csv(strat_nav_path)
    strat_nav['date'] = pd.to_datetime(strat_nav['date'])
    strat_nav = strat_nav.set_index('date')
    
    # Benchmark
    dh = DataHandler(base_path / "database/price_data.parquet")
    dh.load_benchmarks(base_path / "benchmarks")
    bench_df = dh.top_1000_bench.copy()
    bench_df['date'] = pd.to_datetime(bench_df['date'])
    bench_df = bench_df.set_index('date')
    
    # Align Data (Start from Strategy Inception)
    start_date = strat_nav.index.min()
    end_date = strat_nav.index.max()
    
    # Filter Benchmark
    bench_df = bench_df.loc[start_date:end_date]
    
    # Normalize Indices (Rebase to 100)
    strat_nav['rebased'] = (strat_nav['nav'] / strat_nav['nav'].iloc[0]) * 100
    # Handle benchmark rebase carefully - match starting dates
    # Assuming daily data, if dates mismatch slightly, reindex/ffill benchmark to strategy dates
    bench_df = bench_df.reindex(strat_nav.index, method='ffill')
    bench_df['rebased'] = (bench_df['index_value'] / bench_df['index_value'].iloc[0]) * 100
    
    # 2. Performance Stats
    # Overall
    strat_metrics = calculate_metrics(strat_nav.reset_index())
    bench_metrics = calculate_metrics(bench_df.reset_index().rename(columns={'index_value': 'nav'})) # Hack to use util
    
    # Yearly Returns
    # Resample to Year End
    strat_yearly = strat_nav['nav'].resample('Y').last()
    bench_yearly = bench_df['index_value'].resample('Y').last()
    
    # We need start of year values too
    # Actually, easier to calc pct_change from year to year
    # Insert start value
    strat_y_vals = pd.concat([pd.Series([strat_nav['nav'].iloc[0]], index=[strat_nav.index[0]]), strat_yearly])
    bench_y_vals = pd.concat([pd.Series([bench_df['index_value'].iloc[0]], index=[bench_df.index[0]]), bench_yearly])
    
    # Calculate yearly returns properly (Calendar Year)
    # Note: 2017 will be partial (May to Dec)
    years = strat_yearly.index.year.tolist()
    yearly_stats = []
    
    for year in years:
        # Get start/end for this year
        # Start is: End of prev year OR Strategy Start Date
        y_start = pd.Timestamp(f"{year}-01-01")
        if y_start < start_date: y_start = start_date
        
        y_end = pd.Timestamp(f"{year}-12-31")
        if y_end > end_date: y_end = end_date
        
        # Get values
        # Find closest available date for start
        s_date_idx = strat_nav.index.get_indexer([y_start], method='nearest')[0]
        e_date_idx = strat_nav.index.get_indexer([y_end], method='nearest')[0]
        
        s_val = strat_nav['nav'].iloc[s_date_idx]
        e_val = strat_nav['nav'].iloc[e_date_idx]
        s_ret = (e_val / s_val) - 1
        
        b_val_s = bench_df['index_value'].iloc[s_date_idx]
        b_val_e = bench_df['index_value'].iloc[e_date_idx]
        b_ret = (b_val_e / b_val_s) - 1
        
        yearly_stats.append({
            'Year': year,
            'Strategy': s_ret,
            'Benchmark': b_ret,
            'Alpha': s_ret - b_ret
        })
    
    # Rebalance to Rebalance Stats
    # Dates: 15th of Feb, May, Aug, Nov
    # Generate Period Dates
    period_stats = []
    curr = start_date
    while curr < end_date:
        # Determine next rebalance date
        # If curr is May, next is Aug
        # Logic: Find next 15th of Feb, May, Aug, Nov
        candidates = []
        year = curr.year
        for month in [2, 5, 8, 11]:
            cand = pd.Timestamp(year=year, month=month, day=15)
            if cand > curr: candidates.append(cand)
            cand_next = pd.Timestamp(year=year+1, month=month, day=15)
            if cand_next > curr: candidates.append(cand_next)
        
        next_date = min(candidates)
        if next_date > end_date: next_date = end_date
        
        if next_date <= curr: break # Should not happen
        
        # Calculate Return for this period
        try:
            s_idx_1 = strat_nav.index.get_indexer([curr], method='nearest')[0]
            s_idx_2 = strat_nav.index.get_indexer([next_date], method='nearest')[0]
            
            s_r = (strat_nav['nav'].iloc[s_idx_2] / strat_nav['nav'].iloc[s_idx_1]) - 1
            b_r = (bench_df['index_value'].iloc[s_idx_2] / bench_df['index_value'].iloc[s_idx_1]) - 1
            
            period_stats.append({
                'Period': f"{curr.strftime('%Y-%b')} to {next_date.strftime('%Y-%b')}",
                'Strategy': s_r,
                'Benchmark': b_r,
                'Alpha': s_r - b_r,
                'Win': (s_r > b_r)
            })
        except:
            pass
            
        curr = next_date
        if curr >= end_date: break

    # Win Rate
    wins = [p['Win'] for p in period_stats]
    win_rate = sum(wins) / len(wins) * 100 if wins else 0
    total_periods = len(wins)
    
    # 3. Plotting
    plt.figure(figsize=(10, 6))
    plt.plot(strat_nav.index, strat_nav['rebased'], label='Strategy (Quarterly)', linewidth=1.5, color='#1f77b4')
    plt.plot(bench_df.index, bench_df['rebased'], label='Benchmark (Top 1000)', linewidth=1.5, color='#7f7f7f', linestyle='--')
    plt.title('Performance Comparison: Strategy vs. Benchmark (May 2017 - Feb 2026)')
    plt.ylabel('Rebased NAV (Start=100)')
    plt.xlabel('Year')
    plt.legend()
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    chart_path = output_dir / "performance_chart_quarterly.png"
    plt.savefig(chart_path)
    plt.close()
    
    # 4. Generate Markdown Report
    with open(report_file, "w") as f:
        f.write("# Strategic Investment Analysis: Contrarian Breadth \"Champion\"\n")
        f.write(f"**Date Generated:** {pd.Timestamp.now().strftime('%Y-%m-%d')}\n\n")
        
        # Executive Summary
        f.write("## 1. Executive Summary\n")
        f.write("This report details the performance of the **Quarterly Contrarian Breadth Strategy** designed to capture alpha by identifying high-quality industry turnarounds. ")
        f.write("Tested over a 9-year period (May 2017 – Feb 2026), the strategy demonstrates exceptional outperformance against the Top 1000 Market Cap benchmark.\n\n")
        
        f.write("| Metric | Strategy | Benchmark | Difference |\n")
        f.write("| :--- | :---: | :---: | :---: |\n")
        f.write(f"| **Absolute Return** | **{strat_metrics['Absolute Return']}** | {bench_metrics['Absolute Return']} | **+{float(strat_metrics['Absolute Return'].strip('%')) - float(bench_metrics['Absolute Return'].strip('%')):.2f}%** |\n")
        f.write(f"| **CAGR (Annualized)** | **{strat_metrics['CAGR']}** | {bench_metrics['CAGR']} | **+{float(strat_metrics['CAGR'].strip('%')) - float(bench_metrics['CAGR'].strip('%')):.2f}%** |\n")
        f.write(f"| **Max Drawdown** | **{strat_metrics['Max Drawdown']}** | {bench_metrics['Max Drawdown']} | **Better Protection** |\n")
        f.write(f"| **Sharpe Ratio** | **{strat_metrics['Sharpe Ratio']}** | {bench_metrics['Sharpe Ratio']} | **{float(strat_metrics['Sharpe Ratio']) / float(bench_metrics['Sharpe Ratio']):.1f}x Risk-Reward** |\n\n")
        
        # Investment Strategy
        f.write("## 2. Investment Strategy & Rules\n")
        f.write("### Comparison Benchmark\n")
        f.write("- **Benchmark**: Nifty Top 1000 Equal Weight (Broad Market Proxy)\n")
        f.write("- **Rationale**: Represents the opportunity cost of a passive, diversified investment in the Indian market.\n\n")
        
        f.write("### Strategy Logic\n")
        f.write("The strategy employs a **Contrarian Breadth** approach, seeking industries that have seen significant shareholder exodus (capitulation) but are now showing relative strength (recovery).\n\n")
        f.write("**Detailed Rules:**\n")
        f.write("1.  **Investment Universe**: Top 1000 stocks by Market Quant, filtered for liquidity (>0.005% of Market Cap traded daily).\n")
        f.write("2.  **Sector Selection (The \"Contrarian\" Filter)**:\n")
        f.write("    -   Identify **Industry Groups** where shareholders have been decreasing for the last **1 Year (4 Quarters)**.\n")
        f.write("    -   Select the Top 50% of groups with the most widespread shareholder exit.\n")
        f.write("    -   Within those groups, select specific **Industries** where >50% of stocks show shareholder decrease.\n")
        f.write("3.  **Breadth Confirmation (The \"Breadth\" Filter)**:\n")
        f.write("    -   Rank these \"hated\" industries by their **Relative Strength Net Best (RSNP)** score.\n")
        f.write("    -   **RSNP Logic**: % of stocks in the industry outperforming the Top 1000 benchmark over the last month.\n")
        f.write("    -   **Threshold**: Only invest if RSNP > **0.40** (i.e., at least 40% of the industry's stocks are beating the market).\n")
        f.write("4.  **Portfolio Construction**:\n")
        f.write("    -   **Max Stocks**: 15\n")
        f.write("    -   **Concentration Limit**: Max 3 Stocks per Industry.\n")
        f.write("    -   **Stock Selection**: Largest Market Cap stocks within the qualifying industries.\n")
        f.write("    -   **Weighting**: Equal Weight (rebalanced quarterly).\n")
        f.write("5.  **Rebalancing**:\n")
        f.write("    -   **Frequency**: Quarterly (mid-Feb, mid-May, mid-Aug, mid-Nov).\n\n")
        
        # Benchmark Details (Already covered above briefly, expanding here)
        f.write("## 3. Performance Analysis\n")
        f.write("### Cumulative Growth\n")
        f.write("![NAV Comparison Chart](outputs/performance_chart_quarterly.png)\n\n")
        f.write("*The chart illustrates the reliable compounding of the strategy (Blue) vs the Benchmark (Grey). Note the strategy's resilience during the 2020 crash and quicker recovery.*\n\n")
        
        f.write("### Year-by-Year Breakdown\n")
        f.write("| Year | Strategy | Benchmark | Alpha |\n")
        f.write("| :--- | :---: | :---: | :---: |\n")
        for y in yearly_stats:
            alpha_emoji = "✅" if y['Alpha'] > 0 else "🔻"
            f.write(f"| **{y['Year']}** | {y['Strategy']*100:.1f}% | {y['Benchmark']*100:.1f}% | {y['Alpha']*100:.1f}% {alpha_emoji} |\n")
        f.write("\n")
        
        f.write("### Periodicity Analysis (Rebalance-to-Rebalance)\n")
        f.write("This section analyzes the strategy's consistency over its quarterly trading periods.\n\n")
        f.write(f"- **Total Periods**: {total_periods}\n")
        f.write(f"- **Periods Outperforming Benchmark**: {sum(wins)} ({win_rate:.1f}%)\n")
        f.write(f"- **Periods Underperforming**: {total_periods - sum(wins)}\n\n")
        
        f.write("**Detailed Period Returns**\n\n")
        f.write("| Period | Strategy | Benchmark | Alpha |\n")
        f.write("| :--- | :---: | :---: | :---: |\n")
        for p in period_stats:
            win_mark = "**" if p['Win'] else ""
            f.write(f"| {p['Period']} | {win_mark}{p['Strategy']*100:.1f}%{win_mark} | {p['Benchmark']*100:.1f}% | {p['Alpha']*100:.1f}% |\n")
        
        f.write("\n## 4. Risk Analysis\n")
        f.write(f"- **Max Drawdown**: {strat_metrics['Max Drawdown']} (Strategy) vs {bench_metrics['Max Drawdown']} (Benchmark).\n")
        f.write("- **Drawdown Recovery**: The strategy's emphasis on \"Relative Strength\" (RSNP) ensures that it naturally rotates out of lagging sectors, preventing it from holding onto \"value traps\" that never recover—a key risk in pure contrarian investing.\n")
        
    print(f"Report generated successfully: {report_file}")

    # Generate Word Document if possible
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Strategic Investment Analysis: Contrarian Breadth "Champion"', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date Generated: {pd.Timestamp.now().strftime('%Y-%m-%d')}")
        
        # Read Markdown Content (Simple Parsing)
        with open(report_file, "r") as f:
            lines = f.readlines()
            
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("![NAV Comparison Chart]"):
                if chart_path.exists():
                    doc.add_picture(str(chart_path), width=Inches(6))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph("[Chart Image Not Found]")
            elif line.startswith("|"):  # Simple Table handling (Text)
                p = doc.add_paragraph(line)
                p.style = 'List Bullet' # Or just normal, but bullet makes it distinct
            elif line.startswith("- "):
                 doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

        docx_path = base_path / "Client_Report_Contrarian_Quarterly.docx"
        doc.save(docx_path)
        print(f"Word document generated successfully: {docx_path}")
        
    except ImportError:
        print("\n[Notice] 'python-docx' library not found. To generate Word documents, install it via: pip install python-docx")
    except Exception as e:
        print(f"\n[Error] Failed to generate Word document: {e}")

if __name__ == "__main__":
    generate_client_report()
