from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from pathlib import Path
import os

def create_word_report():
    base_path = Path("/Users/shubhrakasana/Desktop/oleander/olgo_mobile/repos/Backtest/precision-backtest-engine")
    report_md_path = base_path / "Client_Report_Contrarian_Quarterly.md"
    chart_path = base_path / "outputs/performance_chart_quarterly.png"
    output_docx_path = base_path / "Client_Report_Contrarian_Quarterly.docx"

    document = Document()

    # Title
    title = document.add_heading('Strategic Investment Analysis: "Risk-Managed Champion"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph(f"Date Generated: {pd.Timestamp.now().strftime('%Y-%m-%d')}")
    
    # Read Markdown Content
    with open(report_md_path, "r") as f:
        lines = f.readlines()
        
    table_mode = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # Handle Tables
        if line.startswith("|"):
            table_mode = True
            row_data = [cell.strip() for cell in line.split("|") if cell.strip()]
            if ":---" in line or "---:" in line: continue # Skip separator row
            table_data.append(row_data)
            continue
        elif table_mode:
            # Table ended, render it
            if table_data:
                table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, val in enumerate(row):
                        if j < len(table.columns):
                            table.cell(i, j).text = val
            table_mode = False
            table_data = []

        if not line: continue

        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("!["):
            # Image handling logic (simplified)
            if "Chart" in line and chart_path.exists():
                document.add_picture(str(chart_path), width=Inches(6))
        else:
            document.add_paragraph(line)

    # Save
    document.save(output_docx_path)
    print(f"Word document saved to: {output_docx_path}")

if __name__ == "__main__":
    create_word_report()
